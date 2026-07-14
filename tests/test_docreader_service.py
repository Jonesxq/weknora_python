from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as WordDocument
from openpyxl import Workbook

from app.api.dependencies import get_document_reader
from app.docreader import (
    DocumentParseError,
    DocumentReaderService,
    UnsupportedDocumentTypeError,
)


@pytest.fixture(scope="module")
def reader() -> DocumentReaderService:
    return DocumentReaderService()


def _minimal_text_pdf(text: str) -> bytes:
    """生成只包含一页文本的最小 PDF，避免测试依赖额外 PDF 生成库。"""

    escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 18 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    data = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(data))
        data.extend(f"{number} 0 obj\n".encode("ascii"))
        data.extend(obj)
        data.extend(b"\nendobj\n")

    xref_offset = len(data)
    data.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    data.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        data.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    data.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(data)


def test_parse_markdown(reader: DocumentReaderService) -> None:
    result = reader.parse_bytes(
        file_name="example.md",
        content="# 标题\n\n这是 Markdown 正文。".encode(),
    )

    assert "标题" in result.content
    assert "Markdown 正文" in result.content
    assert result.images == {}


def test_parse_pdf_text_layer(reader: DocumentReaderService) -> None:
    result = reader.parse_bytes(
        file_name="example.pdf",
        content=_minimal_text_pdf("Hello PDF"),
    )

    assert "Hello PDF" in result.content
    assert result.metadata["page_count"] == 1
    assert result.metadata["text_page_count"] == 1


def test_parse_docx(reader: DocumentReaderService) -> None:
    document = WordDocument()
    document.add_heading("DOCX 标题", level=1)
    document.add_paragraph("DOCX 正文内容")
    content = BytesIO()
    document.save(content)

    result = reader.parse_bytes(
        file_name="example.docx",
        content=content.getvalue(),
    )

    assert "DOCX 标题" in result.content
    assert "DOCX 正文内容" in result.content


def test_parse_xlsx(reader: DocumentReaderService) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "成绩"
    sheet.append(["姓名", "分数"])
    sheet.append(["Alice", 95])
    content = BytesIO()
    workbook.save(content)

    result = reader.parse_bytes(
        file_name="example.xlsx",
        content=content.getvalue(),
    )

    assert "Alice" in result.content
    assert "95" in result.content


def test_parse_path(reader: DocumentReaderService, monkeypatch) -> None:
    # 避免测试依赖操作系统临时目录权限，只验证路径读取和解析编排。
    monkeypatch.setattr(Path, "read_bytes", lambda _path: "# 从路径解析".encode())

    result = reader.parse_path("sample.md")

    assert "从路径解析" in result.content


def test_reject_unsupported_type(reader: DocumentReaderService) -> None:
    with pytest.raises(UnsupportedDocumentTypeError, match="Unsupported file type"):
        reader.parse_bytes(file_name="example.txt", content=b"plain text")


def test_report_corrupt_supported_file_as_parse_error(
    reader: DocumentReaderService,
) -> None:
    with pytest.raises(DocumentParseError, match="解析文档") as exc_info:
        reader.parse_bytes(file_name="broken.xlsx", content=b"not an excel file")

    assert not isinstance(exc_info.value, UnsupportedDocumentTypeError)


def test_dependency_reuses_reader_instance() -> None:
    assert get_document_reader() is get_document_reader()
